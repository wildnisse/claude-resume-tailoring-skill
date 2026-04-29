#!/usr/bin/env python3
"""
Consolidated resume builder. Reads a `resume-content.json` and produces a `.docx`.

Replaces the per-application Python scripts. The structure of resume-content.json
is defined in `schemas/resume-content.schema.json`.

Usage:
    python resume_builder.py --content path/to/resume-content.json --output path/to/output.docx

Optional:
    --style path/to/style.json   override default style settings
"""

from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Emu, Inches, Pt


DEFAULT_STYLE = {
    "font": "Garamond",
    "font_size_pt": 10,
    "name_size_pt": 22,
    "contact_size_pt": 9,
    "section_header_size_pt": 11,
    "summary_shading": "EDEDED",
    "highlights_column_header_smallcaps": True,
    "page_width_emu": 7772400,
    "page_height_emu": 10058400,
    "top_margin_emu": 409575,
    "bottom_margin_emu": 914400,
    "left_margin_in": 0.75,
    "right_margin_in": 0.75,
}


def remove_table_borders(table) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        "</w:tblBorders>"
    )
    tblPr.append(borders)


def add_shading(cell, color: str = "EDEDED") -> None:
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_simple_bullet(cell, text: str, font_name: str, font_size_pt: int) -> None:
    p = cell.add_paragraph()
    run = p.add_run(f"■ {text}")
    run.font.name = font_name
    run.font.size = Pt(font_size_pt)
    pPr = p._p.get_or_add_pPr()
    ind = parse_xml(f'<w:ind {nsdecls("w")} w:left="360" w:hanging="180"/>')
    pPr.append(ind)
    sp = parse_xml(
        f'<w:spacing {nsdecls("w")} w:after="40" w:before="0" w:line="240" w:lineRule="auto"/>'
    )
    existing = pPr.find(qn("w:spacing"))
    if existing is not None:
        pPr.remove(existing)
    pPr.append(sp)


def add_col_header(cell, text: str, font_name: str, font_size_pt: int, smallcaps: bool) -> None:
    cell.paragraphs[0].clear()
    run = cell.paragraphs[0].add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size_pt)
    run.font.bold = True
    if smallcaps:
        rPr = run._r.get_or_add_rPr()
        caps = parse_xml(f'<w:smallCaps {nsdecls("w")} w:val="true"/>')
        rPr.append(caps)


def add_role(
    cell,
    dates: str,
    company_location: str,
    title: str,
    bullets: list[str],
    font_name: str,
    font_size_pt: int,
    subsection_title: str | None = None,
    subsection_bullets: list[str] | None = None,
    first: bool = True,
) -> None:
    if not first:
        sp_p = cell.add_paragraph()
        sp_p.paragraph_format.space_after = Pt(6)
        sp_p.paragraph_format.space_before = Pt(6)

    p_dates = cell.add_paragraph() if not first else cell.paragraphs[0]
    p_dates.clear()
    run_d = p_dates.add_run(dates)
    run_d.font.name = font_name
    run_d.font.size = Pt(font_size_pt)
    run_d.font.bold = True
    pPr_d = p_dates._p.get_or_add_pPr()
    sp_d = parse_xml(
        f'<w:spacing {nsdecls("w")} w:after="0" w:before="0" w:line="240" w:lineRule="auto"/>'
    )
    existing = pPr_d.find(qn("w:spacing"))
    if existing is not None:
        pPr_d.remove(existing)
    pPr_d.append(sp_d)

    p_co = cell.add_paragraph()
    run_co = p_co.add_run(company_location)
    run_co.font.name = font_name
    run_co.font.size = Pt(font_size_pt)
    pPr_co = p_co._p.get_or_add_pPr()
    sp_co = parse_xml(
        f'<w:spacing {nsdecls("w")} w:after="0" w:before="0" w:line="240" w:lineRule="auto"/>'
    )
    existing = pPr_co.find(qn("w:spacing"))
    if existing is not None:
        pPr_co.remove(existing)
    pPr_co.append(sp_co)

    p_title = cell.add_paragraph()
    run_t = p_title.add_run(title)
    run_t.font.name = font_name
    run_t.font.size = Pt(font_size_pt)
    run_t.font.italic = True
    pPr_t = p_title._p.get_or_add_pPr()
    sp_t = parse_xml(
        f'<w:spacing {nsdecls("w")} w:after="40" w:before="0" w:line="240" w:lineRule="auto"/>'
    )
    existing = pPr_t.find(qn("w:spacing"))
    if existing is not None:
        pPr_t.remove(existing)
    pPr_t.append(sp_t)

    for b in bullets:
        add_simple_bullet(cell, b, font_name, font_size_pt)

    if subsection_title:
        p_sub = cell.add_paragraph()
        run_sub = p_sub.add_run(subsection_title)
        run_sub.font.name = font_name
        run_sub.font.size = Pt(font_size_pt)
        run_sub.font.italic = True
        pPr_sub = p_sub._p.get_or_add_pPr()
        sp_sub = parse_xml(
            f'<w:spacing {nsdecls("w")} w:after="40" w:before="60" w:line="240" w:lineRule="auto"/>'
        )
        existing = pPr_sub.find(qn("w:spacing"))
        if existing is not None:
            pPr_sub.remove(existing)
        pPr_sub.append(sp_sub)
        if subsection_bullets:
            for sb in subsection_bullets:
                add_simple_bullet(cell, sb, font_name, font_size_pt)


def build_resume(content: dict[str, Any], style: dict[str, Any], output_path: Path) -> None:
    doc = Document()
    base_style = doc.styles["Normal"]
    base_style.font.name = style["font"]
    base_style.font.size = Pt(style["font_size_pt"])

    section = doc.sections[0]
    section.top_margin = Emu(style["top_margin_emu"])
    section.bottom_margin = Emu(style["bottom_margin_emu"])
    section.left_margin = Inches(style["left_margin_in"])
    section.right_margin = Inches(style["right_margin_in"])
    section.page_width = Emu(style["page_width_emu"])
    section.page_height = Emu(style["page_height_emu"])

    # ── HEADER ────────────────────────────────────────────────────────────────
    header = content["header"]
    contact_parts = []
    for key in ("phone", "email", "linkedin"):
        if header.get(key):
            contact_parts.append(header[key])
    contact_line = "  •  ".join(contact_parts) if contact_parts else ""

    header_table = doc.add_table(rows=2, cols=1)
    remove_table_borders(header_table)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    cell_name = header_table.cell(0, 0)
    cell_name.paragraphs[0].clear()
    run_name = cell_name.paragraphs[0].add_run(header["name"].upper())
    run_name.font.name = style["font"]
    run_name.font.size = Pt(style["name_size_pt"])
    cell_name.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if contact_line:
        p_contact = cell_name.add_paragraph()
        run_contact = p_contact.add_run(contact_line)
        run_contact.font.name = style["font"]
        run_contact.font.size = Pt(style["contact_size_pt"])
        p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER

    cell_loc = header_table.cell(1, 0)
    cell_loc.paragraphs[0].clear()
    if header.get("location"):
        run_loc = cell_loc.paragraphs[0].add_run(header["location"])
        run_loc.font.name = style["font"]
        run_loc.font.size = Pt(style["contact_size_pt"])
        cell_loc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── SUMMARY ───────────────────────────────────────────────────────────────
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    summary_table = doc.add_table(rows=1, cols=3)
    remove_table_borders(summary_table)
    summary_table.columns[0].width = Inches(0.15)
    summary_table.columns[1].width = Inches(6.7)
    summary_table.columns[2].width = Inches(0.15)

    summary_cell = summary_table.cell(0, 1)
    add_shading(summary_cell, style["summary_shading"])
    summary_cell.paragraphs[0].clear()
    run_s = summary_cell.paragraphs[0].add_run(content["summary"])
    run_s.font.name = style["font"]
    run_s.font.size = Pt(style["font_size_pt"])
    pPr = summary_cell.paragraphs[0]._p.get_or_add_pPr()
    sp = parse_xml(
        f'<w:spacing {nsdecls("w")} w:after="0" w:before="0" w:line="240" w:lineRule="auto"/>'
    )
    pPr.append(sp)

    # ── HIGHLIGHTS (3 columns) ────────────────────────────────────────────────
    columns = content.get("highlights", {}).get("columns", [])
    if columns:
        if len(columns) != 3:
            raise ValueError(
                f"Resume content highlights must have exactly 3 columns, got {len(columns)}"
            )

        doc.add_paragraph().paragraph_format.space_after = Pt(2)

        highlights_table = doc.add_table(rows=1, cols=4)
        remove_table_borders(highlights_table)
        highlights_table.columns[0].width = Inches(0.2)
        highlights_table.columns[1].width = Inches(2.2)
        highlights_table.columns[2].width = Inches(2.2)
        highlights_table.columns[3].width = Inches(2.4)

        for i, col in enumerate(columns):
            cell = highlights_table.cell(0, i + 1)
            add_col_header(
                cell,
                col["header"],
                style["font"],
                style["font_size_pt"],
                style["highlights_column_header_smallcaps"],
            )
            for bullet in col["bullets"]:
                add_simple_bullet(cell, bullet, style["font"], style["font_size_pt"])

    # ── PROFESSIONAL EXPERIENCE ───────────────────────────────────────────────
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    exp_header = doc.add_paragraph()
    run_eh = exp_header.add_run("PROFESSIONAL EXPERIENCE")
    run_eh.font.name = style["font"]
    run_eh.font.size = Pt(style["section_header_size_pt"])
    run_eh.font.bold = True
    pPr_eh = exp_header._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="000000"/>'
        "</w:pBdr>"
    )
    pPr_eh.append(pBdr)
    sp_eh = parse_xml(f'<w:spacing {nsdecls("w")} w:after="80" w:before="0"/>')
    pPr_eh.append(sp_eh)

    exp_table = doc.add_table(rows=1, cols=2)
    remove_table_borders(exp_table)
    exp_table.columns[0].width = Inches(0.25)
    exp_table.columns[1].width = Inches(6.75)
    exp_cell = exp_table.cell(0, 1)
    exp_cell.paragraphs[0].clear()

    for i, role in enumerate(content["experience"]):
        sub = role.get("subsection")
        add_role(
            exp_cell,
            role["dates"],
            role["company_location"],
            role["title"],
            role["bullets"],
            style["font"],
            style["font_size_pt"],
            subsection_title=sub["title"] if sub else None,
            subsection_bullets=sub["bullets"] if sub else None,
            first=(i == 0),
        )

    # ── EDUCATION ─────────────────────────────────────────────────────────────
    if content.get("education"):
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

        edu_header = doc.add_paragraph()
        run_edu_h = edu_header.add_run("EDUCATION")
        run_edu_h.font.name = style["font"]
        run_edu_h.font.size = Pt(style["section_header_size_pt"])
        run_edu_h.font.bold = True
        pPr_edu = edu_header._p.get_or_add_pPr()
        pBdr_edu = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            '  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="000000"/>'
            "</w:pBdr>"
        )
        pPr_edu.append(pBdr_edu)
        sp_edu = parse_xml(f'<w:spacing {nsdecls("w")} w:after="80" w:before="0"/>')
        pPr_edu.append(sp_edu)

        edu_table = doc.add_table(rows=1, cols=2)
        remove_table_borders(edu_table)
        edu_table.columns[0].width = Inches(0.25)
        edu_table.columns[1].width = Inches(6.75)
        edu_cell = edu_table.cell(0, 1)
        edu_cell.paragraphs[0].clear()

        for i, edu in enumerate(content["education"]):
            if i == 0:
                p_inst = edu_cell.paragraphs[0]
            else:
                p_inst = edu_cell.add_paragraph()
            run_inst = p_inst.add_run(edu["institution"])
            run_inst.font.name = style["font"]
            run_inst.font.size = Pt(style["font_size_pt"])

            p_deg = edu_cell.add_paragraph()
            run_deg = p_deg.add_run(edu["degree"])
            run_deg.font.name = style["font"]
            run_deg.font.size = Pt(style["font_size_pt"])
            run_deg.font.italic = True

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def main() -> int:
    parser = argparse.ArgumentParser(description="Build a resume docx from resume-content.json")
    parser.add_argument(
        "--content", required=True, type=Path, help="Path to resume-content.json"
    )
    parser.add_argument("--output", required=True, type=Path, help="Output docx path")
    parser.add_argument(
        "--style", type=Path, default=None, help="Optional style.json overriding defaults"
    )
    args = parser.parse_args()

    if not args.content.exists():
        print(f"ERROR: content file not found: {args.content}", file=sys.stderr)
        return 1

    content = json.loads(args.content.read_text())

    style = dict(DEFAULT_STYLE)
    if "style" in content and content["style"]:
        style.update(content["style"])
    if args.style:
        if not args.style.exists():
            print(f"ERROR: style file not found: {args.style}", file=sys.stderr)
            return 1
        style.update(json.loads(args.style.read_text()))

    build_resume(content, style, args.output)
    print(f"Saved: {args.output}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
